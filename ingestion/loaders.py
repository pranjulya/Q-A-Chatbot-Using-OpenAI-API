"""Document loading utilities."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Sequence

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}


@dataclass
class Document:
    """Simple representation of a source document."""

    text: str
    path: Path

    @property
    def metadata(self) -> dict[str, str]:
        return {
            "source": str(self.path),
            "name": self.path.name,
        }


def iter_files(root: Path) -> Iterable[Path]:
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path


def load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError(
            "pdfplumber is required to parse PDF files. Install it via 'pip install pdfplumber'."
        ) from exc

    with pdfplumber.open(path) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages)


def load_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError(
            "python-docx is required to parse DOCX files. Install it via 'pip install python-docx'."
        ) from exc

    document = docx.Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


LOADERS = {
    ".txt": load_text,
    ".md": load_text,
    ".markdown": load_text,
    ".pdf": load_pdf,
    ".docx": load_docx,
}


def load_documents(paths: Sequence[Path]) -> List[Document]:
    documents: List[Document] = []
    for path in paths:
        suffix = path.suffix.lower()
        loader = LOADERS.get(suffix)
        if not loader:
            continue
        text = loader(path)
        if text.strip():
            documents.append(Document(text=text, path=path))
    return documents


def load_documents_from_directory(directory: Path) -> List[Document]:
    file_paths = list(iter_files(directory))
    if not file_paths:
        raise FileNotFoundError(
            f"No supported documents found in {directory}. Supported extensions: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    return load_documents(file_paths)
